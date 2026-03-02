import io
import os
import logging
from typing import Tuple, Optional
import pypdf
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Utility class for processing PDF and Word documents"""
    
    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
    MAX_TEXT_LENGTH = 100000  # Limit extracted text to prevent token overflow
    
    @staticmethod
    def extract_text_from_pdf(file_data: bytes) -> Tuple[bool, str]:
        """
        Extract text from PDF file data
        
        Args:
            file_data: Raw PDF file bytes
            
        Returns:
            Tuple of (success: bool, text: str or error_message: str)
        """
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                return False, "PDF file appears to be empty or corrupted"
            
            buf = io.StringIO()
            max_pages = 100  # Limit to first 100 pages to prevent excessive processing
            truncated = False
            
            for page_num, page in enumerate(pdf_reader.pages[:max_pages]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        buf.write(f"\n--- Page {page_num + 1} ---\n")
                        buf.write(page_text)
                        
                        # Stop if we've extracted enough text
                        if buf.tell() > DocumentProcessor.MAX_TEXT_LENGTH:
                            truncated = True
                            break
                            
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            text_content = buf.getvalue()
            if truncated:
                text_content = text_content[:DocumentProcessor.MAX_TEXT_LENGTH]
                text_content += "\n\n[Document truncated - maximum text length reached]"
            
            if not text_content.strip():
                return False, "No readable text found in PDF. The PDF might contain only images or be password-protected."
            
            return True, text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return False, f"Failed to process PDF file: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file_data: bytes) -> Tuple[bool, str]:
        """
        Extract text from Word document file data
        
        Args:
            file_data: Raw DOCX file bytes
            
        Returns:
            Tuple of (success: bool, text: str or error_message: str)
        """
        try:
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            
            buf = io.StringIO()
            truncated = False
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    buf.write(paragraph.text + "\n")
                    if buf.tell() > DocumentProcessor.MAX_TEXT_LENGTH:
                        truncated = True
                        break
            
            # Extract text from tables if not too much text already
            if not truncated:
                for table in doc.tables:
                    buf.write("\n--- Table ---\n")
                    for row in table.rows:
                        row_text = [
                            cell.text.strip()
                            for cell in row.cells
                            if cell.text.strip()
                        ]
                        if row_text:
                            buf.write(" | ".join(row_text) + "\n")
                        if buf.tell() > DocumentProcessor.MAX_TEXT_LENGTH:
                            truncated = True
                            break
                    if truncated:
                        break
            
            text_content = buf.getvalue()
            if truncated:
                text_content = text_content[:DocumentProcessor.MAX_TEXT_LENGTH]
                text_content += "\n\n[Document truncated - maximum text length reached]"
            
            if not text_content.strip():
                return False, "No readable text found in Word document."
            
            return True, text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            return False, f"Failed to process Word document: {str(e)}"
    
    @staticmethod
    def process_document(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
        """
        Process document based on file extension
        
        Args:
            file_data: Raw file bytes
            filename: Original filename with extension
            
        Returns:
            Tuple of (success: bool, content: str, file_type: str or None)
        """
        if len(file_data) > DocumentProcessor.MAX_FILE_SIZE:
            return False, "File size too large. Please upload a document smaller than 20MB.", None
        
        file_extension = os.path.splitext(filename.lower())[1]
        
        if file_extension == '.pdf':
            success, content = DocumentProcessor.extract_text_from_pdf(file_data)
            return success, content, 'pdf'
        elif file_extension in ['.docx', '.doc']:
            if file_extension == '.doc':
                return False, "Please convert .doc files to .docx format. Only .docx files are supported.", None
            success, content = DocumentProcessor.extract_text_from_docx(file_data)
            return success, content, 'docx'
        else:
            return False, f"Unsupported file type: {file_extension}. Only PDF (.pdf) and Word (.docx) documents are supported.", None
    
    @staticmethod
    def validate_document_file(filename: str, file_size: int) -> Tuple[bool, str]:
        """
        Validate document file before processing
        
        Args:
            filename: Original filename
            file_size: File size in bytes
            
        Returns:
            Tuple of (is_valid: bool, error_message: str or empty string)
        """
        if not filename:
            return False, "No file selected"
        
        if file_size > DocumentProcessor.MAX_FILE_SIZE:
            return False, "File size too large. Please upload a document smaller than 20MB."
        
        file_extension = os.path.splitext(filename.lower())[1]
        supported_extensions = {'.pdf', '.docx'}
        
        if file_extension not in supported_extensions:
            return False, f"Unsupported file type: {file_extension}. Only PDF (.pdf) and Word (.docx) documents are supported."
        
        if file_extension == '.doc':
            return False, "Please convert .doc files to .docx format. Only .docx files are supported."
        
        return True, ""
